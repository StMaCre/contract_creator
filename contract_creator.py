# Import
import os
import datetime
from docx import Document
from docx.shared import Inches 
import google_setup # Assuming google_setup.py is in the same directory
import re

# --- Configuration ---
TEMPLATE_DOCX_PATH = "Contracts to be created\PPMI ENESET service contracts - special conditions_template.docx"       # Path to your Word template
ANNEX1_TEMPLATE_DOCX_PATH = "Contracts to be created\Annex 1 - Technical specification and timeline of the report.docx"
MINUTES_DOCX_PATH = "Contracts to be created\ENESET kick off AR Successful learning trajectories minutes _v1.docx"        # Path to the text file with OCR'd minutes
OUTPUT_DIR = "generated_contracts"        # Directory to save the filled contracts

# --- Predefined Contract Details (User Input) ---
# You would typically get these from user input or another system
CONTRACT_NUMBER_BASE = "2024-DG EAC-ENESET" # Base part, number gets added
NEXT_CONTRACT_SEQ = 13 # The sequential number for this specific contract
CONTRACT_DATE = datetime.date.today().strftime("%Y-%m-%d") # Or a specific date string
DELIVERABLE_COORDINATOR = "Stéphanie Crêteur" # Example coordinator name


# --- Placeholders (MUST map keys used in data dicts to {{placeholders}} in TEMPLATES) ---
# Contains placeholders potentially used in EITHER template.
# The fill function will only replace placeholders for which data is provided in the specific call.
PLACEHOLDERS = {
    # For Main Contract
    "contract_no": "{{CONTRACT_NUMBER}}",
    "contract_date": "{{CONTRACT_DATE}}",
    "report_no": "{{REPORT_NUMBER}}",
    "deliverable_coordinator": "{{DELIVERABLE_COORDINATOR}}",

    # For BOTH Main Contract AND Annex 1 Template (ensure these are in annex1_template.docx)
    "report_name": "{{REPORT_NAME}}",
    "report_objective": "{{REPORT_OBJECTIVE}}",

    # For Annex 1 Template ONLY (ensure this is in annex1_template.docx)
    "timeline_summary": "{{TIMELINE_SUMMARY}}",

    # Removed Annex-specific placeholders like {{ANNEX_REPORT_NAME}} as they are not needed
    # when using separate templates with the direct {{REPORT_NAME}} etc. placeholders.
}

# --- LLM Interaction (with Prompt Updates) ---

def extract_info_from_minutes(minutes_text, generate_text_func):
    """Extracts report details, requests British English, longer objectives,
       and applies specific text replacements."""
    print("Extracting information from minutes text using LLM...")
    extracted_data = {}

    # --- Updated Prompts ---
    prompts = {
        "report_no": (
            "From the following meeting minutes text, extract the specific Analytical Report number "
            "(like 'AR3').\n\nText:\n---\n" f"{minutes_text}\n---\nReport Number:"
        ),
        "report_name": (
            "From the following meeting minutes text, extract the full official title/name of the report. "
            "Ensure correct capitalisation. Use British English spelling if applicable.\n\n"
            "Text:\n---\n" f"{minutes_text}\n---\nReport Name:"
        ),
        # Prompt for longer objective and British English
        "report_objective": (
            "Based on the 'Overview of the request' or 'teaser' box in the following minutes text, "
            "generate a detailed description of the report's purpose and objectives, suitable for a formal contract. "
            "Start with a brief introductory sentence summarising the report's main focus. Then, elaborate on the specific "
            "objectives mentioned (like the three points in the teaser: examining barriers, analysing replicability, outlining trajectories). "
            "Ensure the description is comprehensive and clearly outlines the scope.\n\n"
            "**Important:** Please use **British English spelling** throughout (e.g., 'analyse', 'organisation', 'programme').\n\n"
            "Text:\n---\n"
            f"{minutes_text}\n---\n"
            "Detailed Report Objective Description (British English):"
        ),
         "timeline_summary": (
            "From the 'Timeline and next steps' table in the minutes, extract and list the key dates and activities. "
            "Format strictly as 'DD Month YYYY - Activity description', with each entry on a new line.\n\n"
            "Text:\n---\n" f"{minutes_text}\n---\n"
            "Timeline Summary (DD Month YYYY - Activity, one per line):"
         )
    }

    for key, prompt in prompts.items():
        print(f"  > Asking LLM for: {key}")
        # Use a config that allows for longer output, especially for objective
        config = google_setup.genai.types.GenerationConfig(temperature=0.25, max_output_tokens=1500) # Slightly higher temp, more tokens
        response = generate_text_func(prompt, config=config)
        cleaned_response = response.strip().strip('"`')

        if "[[GENERATION ERROR" in cleaned_response or "[[RESPONSE BLOCKED" in cleaned_response or "[[NO TEXT" in cleaned_response:
             print(f"  > Warning: LLM failed or returned no content for {key}. Response: {cleaned_response}")
             extracted_data[key] = f"[[ERROR EXTRACTING {key.upper()}]]"
        else:
            final_text = cleaned_response
            # --- Post-processing for specific replacements (applied to objective) ---
            if key == "report_objective":
                print("  > Applying 'competence(s)' replacement to objective...")
                # Use regex for case-insensitivity and variations
                # Replace 'competencies', 'competency', 'competencie(s)' with 'competences', 'competence', 'competence(s)'
                final_text = re.sub(r'competencies', 'competences', final_text, flags=re.IGNORECASE)
                final_text = re.sub(r'competency', 'competence', final_text, flags=re.IGNORECASE)
                # Handle the specific "(s)" form if present
                final_text = re.sub(r'competencie\(s\)', 'competence(s)', final_text, flags=re.IGNORECASE)
                print(f"  > Text after replacement: {final_text[:150]}...")

            extracted_data[key] = final_text
            print(f"  > Received for {key} (final): {final_text[:150]}...")

    print("LLM extraction complete.")
    return extracted_data

# --- Word Document Processing (using run-level replacement - unchanged) ---

def replace_text_in_runs(container, placeholder, value):
    """Replaces placeholder text within runs of paragraphs in a container (doc, cell)."""
    for paragraph in container.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, str(value))

def fill_word_template(template_path, output_path, data):
    """Fills placeholders in a Word template using run-level replacement."""
    print(f"Filling template: {template_path}")
    try:
        document = Document(template_path)
        for key, value in data.items():
            placeholder = PLACEHOLDERS.get(key)
            if placeholder:
                replace_text_in_runs(document, placeholder, value)
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            replace_text_in_runs(cell, placeholder, value)
            else:
                 print(f"Warning: No placeholder mapping for data key '{key}'.")
        document.save(output_path)
        print(f"Successfully generated document: {output_path}")
        return True
    except FileNotFoundError: print(f"Error: Template file not found at {template_path}"); return False
    except KeyError as e: print(f"Error: Missing key {e} in PLACEHOLDERS map?"); return False
    except Exception as e: print(f"Error processing Word document '{template_path}': {e}"); return False

# --- Main Execution (structure unchanged) ---
if __name__ == "__main__":
    print("Starting contract and Annex 1 generation process...")

    # 1. Setup LLM Client
    try:
        generate_text = google_setup.setup_genai_client()
        print("Google GenAI client setup successful.")
    except Exception as e: print(f"Fatal Error: Could not set up Google GenAI client. {e}"); exit()

    # 2. Load Minutes from DOCX
    minutes_content = None
    try:
        print(f"Attempting to read minutes from DOCX: {MINUTES_DOCX_PATH}")
        doc = Document(MINUTES_DOCX_PATH)
        all_text = all_text = [para.text for para in doc.paragraphs] # Get text from each paragraph
        for table in doc.tables: # Add tables
            all_text.append("\n--- Table Start ---")
            for row in table.rows:
                row_text = ["\n".join(p.text for p in cell.paragraphs).strip() for cell in row.cells]
                all_text.append(" | ".join(row_text))
            all_text.append("--- Table End ---\n")
        minutes_content = "\n".join(all_text)
        if not minutes_content.strip(): print(f"Warning: Extracted text from {MINUTES_DOCX_PATH} appears empty.")
        print(f"Successfully loaded minutes text from: {MINUTES_DOCX_PATH}")
    except FileNotFoundError: print(f"Fatal Error: Minutes DOCX file not found at {MINUTES_DOCX_PATH}"); exit()
    except Exception as e: print(f"Fatal Error: Could not read minutes DOCX file '{MINUTES_DOCX_PATH}'. Error: {e}"); exit()

    # 3. Extract Info using LLM
    if minutes_content:
        llm_extracted_data = extract_info_from_minutes(minutes_content, generate_text)
    else: print("Fatal Error: Minutes content is empty."); exit()

    # 4. Prepare Data Dictionaries
    full_contract_number = f"{CONTRACT_NUMBER_BASE} / No {NEXT_CONTRACT_SEQ:03d}"
    report_no = llm_extracted_data.get("report_no", "[[Report No. Not Found]]")
    report_name = llm_extracted_data.get("report_name", "[[Report Name Not Found]]")
    # The objective now comes with replacements already applied
    report_objective = llm_extracted_data.get("report_objective", "[[Report Objective Not Found]]")
    timeline_summary = llm_extracted_data.get("timeline_summary", "[[Timeline Not Found]]")

    contract_data = {
        "contract_no": full_contract_number, "contract_date": CONTRACT_DATE,
        "report_no": report_no, "report_name": report_name,
        "report_objective": report_objective, "deliverable_coordinator": DELIVERABLE_COORDINATOR,
    }
    annex1_data = {
        "report_name": report_name, "report_objective": report_objective,
        "timeline_summary": timeline_summary,
    }

    # 5. Create Output Directory
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # 6. Generate Filenames
    report_no_sanitized = report_no.replace(" ", "_").replace("/", "-").replace("\\", "-")
    base_filename = f"Contract_{report_no_sanitized}_{CONTRACT_DATE}"
    contract_output_filename = f"{base_filename}.docx"
    annex1_output_filename = f"Annex1_{base_filename}.docx"
    contract_output_path = os.path.join(OUTPUT_DIR, contract_output_filename)
    annex1_output_path = os.path.join(OUTPUT_DIR, annex1_output_filename)

    # 7. Fill Main Contract Template
    print("\n--- Generating Main Contract ---")
    success_contract = fill_word_template(TEMPLATE_DOCX_PATH, contract_output_path, contract_data)

    # 8. Fill Annex 1 Template
    print("\n--- Generating Annex 1 ---")
    success_annex1 = fill_word_template(ANNEX1_TEMPLATE_DOCX_PATH, annex1_output_path, annex1_data)

    # 9. Final Status
    print("\n--- Summary ---")
    if success_contract: print(f"Main Contract generated: {contract_output_path}")
    else: print(f"Main Contract generation FAILED.")
    if success_annex1: print(f"Annex 1 generated: {annex1_output_path}")
    else: print(f"Annex 1 generation FAILED.")
    if success_contract and success_annex1: print("\nProcess completed.")
    else: print("\nProcess finished with errors.")